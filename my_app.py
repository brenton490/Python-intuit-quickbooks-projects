from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak(
    'Welcome to My CV Builder! Please answer the following questions to build your CV.')


def speak(text):
    pyttsx3.speak(text)


document = Document()
# Picture
document.add_picture(
    'bw.jpg',
    width=Inches(2.0))
# Name phone number email details
speak('What is your name?')
name = input('What is your name?')
speak('Hello' + name + 'nice to meet you!')
speak('What is your phone number?')
phone_number = input('What is your phone number?')
speak('Your phone number is' + phone_number)
speak('What is your email address?')
email = input('What is your email address?')
speak('Your email address is' + email)
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me section
document.add_heading('About Me')
speak('Tell me about yourself')
about_me = input('Tell me about yourself?')
document.add_paragraph(about_me)
# Work Experience section
document.add_heading('Work Experience')

p = document.add_paragraph()
speak('Please enter the company name here:')
company = input('Enter company name: ')
speak('From which date did you work there?')
from_date = input('From Date: ')
speak('Until which date did you work there?')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
speak('Describe your experience at' + company)
experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# more experiences
while True:
    speak('Do you have more work experiences? Yes or No')
    has_more_experiences = input(
        'Do you have more work experiences? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        speak('Please enter the company name here:')
        company = input('Enter company name: ')
        speak('From which date did you work there?')
        from_date = input('From Date: ')
        speak('Until which date did you work there?')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        speak('Describe your experience at' + company)
        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break
# Skills section
document.add_heading('Skills')
speak('What are your skills?')
skill = input("what are your skills? ")
while True:
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'
    speak('Do you have more skills? Yes or No')
    more_skills = input("Do you have more skills? Yes or No: ")
    if more_skills.lower() == 'yes':
        speak('What are your skills?')
        skill = input("what are your skills? ")
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using My CV Builder"
# Save Document
document.save('cv.docx')
